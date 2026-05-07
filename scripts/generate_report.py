"""Daily cloud intelligence report generator.

Calls the Anthropic API with the web_search tool, then converts the resulting
Markdown report to a DOCX file under reports/.
"""

from __future__ import annotations

import os
import re
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

JST = timezone(timedelta(hours=9))
ROOT = Path(__file__).resolve().parent.parent
PROMPT_PATH = ROOT / "prompts" / "daily-cloud-report.md"
REPORTS_DIR = ROOT / "reports"
MODEL = "claude-opus-4-7"


def render_prompt(today: str, yesterday: str) -> str:
    template = PROMPT_PATH.read_text(encoding="utf-8")
    return template.replace("{{TODAY}}", today).replace("{{YESTERDAY}}", yesterday)


def call_claude(prompt: str) -> str:
    client = anthropic.Anthropic()
    parts: list[str] = []

    with client.messages.stream(
        model=MODEL,
        max_tokens=32000,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        tools=[{"type": "web_search_20260209", "name": "web_search"}],
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        for text in stream.text_stream:
            parts.append(text)

        final = stream.get_final_message()

    if not parts:
        for block in final.content:
            if block.type == "text":
                parts.append(block.text)

    markdown = "".join(parts).strip()
    if not markdown:
        raise RuntimeError("Empty response from Claude API")
    return markdown


def md_to_docx(markdown: str, out_path: Path, today: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10)

    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(table_rows[0]):
            table.rows[0].cells[i].text = h
        for row_cells in table_rows[1:]:
            row = table.add_row().cells
            for i, cell in enumerate(row_cells[:cols]):
                row[i].text = cell
        table_rows = []
        in_table = False

    lines = markdown.splitlines()
    for raw in lines:
        line = raw.rstrip()

        if line.startswith("|") and "|" in line[1:]:
            if re.match(r"^\|[\s:|-]+\|\s*$", line):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table()

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        elif line == "":
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)

    flush_table()

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run(f"Generated: {today} (JST)")
    r.font.size = Pt(8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY is not set", file=sys.stderr)
        return 1

    now_jst = datetime.now(JST)
    today = now_jst.strftime("%Y-%m-%d")
    yesterday = (now_jst - timedelta(days=1)).strftime("%Y-%m-%d")
    compact = now_jst.strftime("%Y%m%d")

    print(f"[info] generating report for {today} (JST)")
    prompt = render_prompt(today, yesterday)

    markdown = call_claude(prompt)

    md_path = REPORTS_DIR / f"CloudReport_{compact}.md"
    docx_path = REPORTS_DIR / f"CloudReport_{compact}.docx"

    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(markdown, encoding="utf-8")
    md_to_docx(markdown, docx_path, today)

    print(f"[ok] wrote {md_path.relative_to(ROOT)}")
    print(f"[ok] wrote {docx_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
